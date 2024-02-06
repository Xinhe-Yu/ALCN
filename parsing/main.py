# Import the required library
from docx import Document
import os
import re

def extract_chinese_and_foreign_terms(text):
    """
    Extracts and pairs Chinese terms with their corresponding foreign terms in parentheses.

    Parameters:
    text (str): The input text from which to extract terms.

    Returns:
    list of str: A list of extracted terms in "Chinese;Foreign" format.
    """
    results = []
    pattern = re.compile(r'([\u4e00-\u9fff·\-”》]+)\（([^\（\）]+)\）')

    for match in pattern.finditer(text):
        chinese_context = match.group(1)
        foreign_term = match.group(2)

        # Determine the starting point for capturing Chinese characters based on the foreign term length
        start_index = max(0, len(chinese_context) - len(foreign_term))

        # Extract the relevant Chinese segment
        chinese_segment = chinese_context[start_index:]
        results.append(f"{chinese_segment};{foreign_term}")
    return results

def read_word_file(file_path):
    """
    Reads a Word document and returns its text content.

    Parameters:
        file_path (str): The path to the Word document.

    Returns:
        str: The text content of the Word document.
    """
    # Load the Word document
    doc = Document(file_path)

    # Read and return the text content
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    if hasattr(doc, 'footnotes'):
        print("Reading footnotes...")
        for footnote in doc.footnotes:
            for para in footnote.paragraphs:
                full_text.append(para.text)

    return '\n'.join(full_text)

def write_extracted_terms_to_file(extracted_terms, docx_file_path):
    # Extraire le nom de base du fichier Word pour le nom du fichier de sortie
    base_name = os.path.basename(docx_file_path)

    # Enlever l'extension .docx et ajouter .txt
    output_file_name = os.path.splitext(base_name)[0] + ".txt"
    output_folder_path = "parsing/extracted_texts"
    output_file_path = os.path.join(output_folder_path, output_file_name)

    if not os.path.exists(output_folder_path):
        os.makedirs(output_folder_path)

    with open(output_file_path, "w", encoding="utf-8") as file:
        for term in extracted_terms:
            file.write(f"{term}\n")

    print(f"Les termes extraits ont été écrits dans '{output_file_path}'.")


def read_documents_in_folder(folder_path):
    """
    Reads all Word documents in the specified folder and prints their content.

    Parameters:
        folder_path (str): The path to the folder containing Word documents.
    """
    # Iterate over all files in the folder
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            print(f"Reading {filename} content...")
            content = read_word_file(file_path)
            extracted_terms = extract_chinese_and_foreign_terms(content)
            write_extracted_terms_to_file(extracted_terms, filename)
            print("-" * 40)  # Separator between documents

documents_path = "parsing/documents"

# Read and print the content of all Word documents in the 'documents' folder
read_documents_in_folder(documents_path)
